"""
Service de génération de documents Word à partir de la synthèse de financement
"""

from datetime import datetime
from pathlib import Path
from typing import Dict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import insert_document_genere


class WordDocumentGenerator:
    """Service pour générer des documents Word à partir des données de synthèse"""

    def __init__(self):
        """Initialiser le service de génération de document Word"""
        # Créer le dossier de sortie s'il n'existe pas
        self.output_dir = Path("generated_documents")
        self.output_dir.mkdir(exist_ok=True)
        print("✅ WordDocumentGenerator initialisé")

    def _add_heading_with_style(self, doc: Document, text: str, level: int = 1):
        """
        Ajouter un titre avec un style personnalisé

        Args:
            doc: Document Word
            text: Texte du titre
            level: Niveau du titre (1, 2, 3)
        """
        heading = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Appliquer un style en gras
        for run in heading.runs:
            run.bold = True

    def _add_paragraph_with_style(self, doc: Document, text: str, bold: bool = False):
        """
        Ajouter un paragraphe avec style

        Args:
            doc: Document Word
            text: Texte du paragraphe
            bold: Mettre en gras
        """
        para = doc.add_paragraph(text)
        if bold:
            for run in para.runs:
                run.bold = True
        return para

    def _safe_get(self, data: Dict, key: str, default: str = "Non spécifié") -> str:
        """
        Récupérer une valeur de manière sécurisée

        Args:
            data: Dictionnaire source
            key: Clé à récupérer
            default: Valeur par défaut si la clé n'existe pas

        Returns:
            str: Valeur récupérée ou default
        """
        value = data.get(key, default)
        if value is None or value == "":
            return default
        return str(value)

    def generate_word_document(self, synthese_data: Dict, dossier_id: str) -> Dict:
        """
        Générer un document Word à partir de la synthèse JSON

        Args:
            synthese_data: Données de la synthèse
            dossier_id: Identifiant du dossier

        Returns:
            Dict: Résultat de la génération avec chemin du fichier
        """
        try:
            # Créer le nom du fichier
            file_name = f"Carte_Financement_{dossier_id}.docx"
            file_path = self.output_dir / file_name

            # Créer le document
            doc = Document()

            # === EN-TÊTE ===
            title = doc.add_heading('CARTE DE FINANCEMENT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.bold = True

            # Informations du dossier
            doc.add_paragraph()
            subtitle = doc.add_paragraph(f"Dossier n° : {dossier_id}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in subtitle.runs:
                run.bold = True

            date_gen = doc.add_paragraph(f"Généré le : {synthese_data.get('date_generation', datetime.now().strftime('%d/%m/%Y %H:%M:%S'))}")
            date_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Séparateur
            doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # === 1. SYNTHÈSE DU PROJET ===
            self._add_heading_with_style(doc, "1. SYNTHÈSE DU PROJET", 1)
            synthese_projet = synthese_data.get('synthese_projet', {})

            if synthese_projet:
                doc.add_paragraph(f"Description : {self._safe_get(synthese_projet, 'description')}")
                doc.add_paragraph(f"Objectif : {self._safe_get(synthese_projet, 'objectif_financement')}")
                doc.add_paragraph(f"Lieu : {self._safe_get(synthese_projet, 'lieu')}")
                doc.add_paragraph(f"Montant total : {self._safe_get(synthese_projet, 'montant_total')}")
                doc.add_paragraph(f"Durée : {self._safe_get(synthese_projet, 'duree')}")
                doc.add_paragraph(f"Garanties : {self._safe_get(synthese_projet, 'garanties')}")
            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === 2. PROFIL EMPRUNTEUR ===
            self._add_heading_with_style(doc, "2. PROFIL EMPRUNTEUR", 1)
            profil_emprunteur = synthese_data.get('profil_emprunteur', {})

            if profil_emprunteur:
                # Section Identité
                identite = profil_emprunteur.get('identite', {})
                if identite:
                    self._add_heading_with_style(doc, "Identité", 2)
                    doc.add_paragraph(f"Civilité : {self._safe_get(identite, 'civilite')}")
                    doc.add_paragraph(f"Nom : {self._safe_get(identite, 'nom')}")
                    doc.add_paragraph(f"Prénoms : {self._safe_get(identite, 'prenoms')}")
                    doc.add_paragraph(f"Date de naissance : {self._safe_get(identite, 'date_naissance')}")
                    doc.add_paragraph(f"Lieu de naissance : {self._safe_get(identite, 'lieu_naissance')}")
                    doc.add_paragraph(f"Nationalité : {self._safe_get(identite, 'nationalite')}")
                    doc.add_paragraph(f"Profession : {self._safe_get(identite, 'profession')}")
                    doc.add_paragraph(f"Email : {self._safe_get(identite, 'email')}")
                    doc.add_paragraph(f"Téléphone : {self._safe_get(identite, 'telephone')}")

                # Situation familiale
                situation_familiale = self._safe_get(profil_emprunteur, 'situation_familiale')
                if situation_familiale != "Non spécifié":
                    self._add_heading_with_style(doc, "Situation familiale", 2)
                    doc.add_paragraph(f"Situation : {situation_familiale}")
                    doc.add_paragraph(f"Régime matrimonial : {self._safe_get(profil_emprunteur, 'regime_matrimonial')}")
                    doc.add_paragraph(f"Enfants à charge : {self._safe_get(profil_emprunteur, 'enfants_a_charge')}")

                # Adresse
                adresse = profil_emprunteur.get('adresse')
                if adresse and isinstance(adresse, dict):
                    self._add_heading_with_style(doc, "Adresse", 2)
                    doc.add_paragraph(f"Numéro voie : {self._safe_get(adresse, 'numero_voie')}")
                    doc.add_paragraph(f"Nom voie : {self._safe_get(adresse, 'nom_voie')}")
                    doc.add_paragraph(f"Code postal : {self._safe_get(adresse, 'code_postal')}")
                    doc.add_paragraph(f"Ville : {self._safe_get(adresse, 'ville')}")
                    doc.add_paragraph(f"Pays : {self._safe_get(adresse, 'pays')}")
            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === 3. REVENUS ===
            self._add_heading_with_style(doc, "3. REVENUS", 1)
            revenus = synthese_data.get('revenus', {})

            if revenus:
                doc.add_paragraph(f"Revenus annuels moyens : {self._safe_get(revenus, 'revenus_annuels_moyens')}")
                doc.add_paragraph(f"Dernier revenu fiscal : {self._safe_get(revenus, 'dernier_revenu_fiscal')}")
                doc.add_paragraph(f"Revenus mensuels : {self._safe_get(revenus, 'revenus_mensuels')}")
                doc.add_paragraph(f"Bonus et primes : {self._safe_get(revenus, 'bonus_primes')}")
                doc.add_paragraph(f"Revenus fonciers : {self._safe_get(revenus, 'revenus_fonciers')}")
                doc.add_paragraph(f"Autres revenus : {self._safe_get(revenus, 'autres_revenus')}")
            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === 4. PATRIMOINE IMMOBILIER ===
            self._add_heading_with_style(doc, "4. PATRIMOINE IMMOBILIER", 1)
            patrimoine_immobilier = synthese_data.get('patrimoine_immobilier', {})

            if patrimoine_immobilier:
                doc.add_paragraph(f"Biens immobiliers : {self._safe_get(patrimoine_immobilier, 'biens_immobiliers')}")
                doc.add_paragraph(f"Valeur estimée totale : {self._safe_get(patrimoine_immobilier, 'valeur_estimee_totale')}")
                doc.add_paragraph(f"Crédits restants dus : {self._safe_get(patrimoine_immobilier, 'credits_restants_dus')}")
                doc.add_paragraph(f"Loyers perçus annuels : {self._safe_get(patrimoine_immobilier, 'loyers_percus_annuels')}")
                doc.add_paragraph(f"Patrimoine net immobilier : {self._safe_get(patrimoine_immobilier, 'patrimoine_net_immobilier')}")
            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === 5. PATRIMOINE MOBILIER ===
            self._add_heading_with_style(doc, "5. PATRIMOINE MOBILIER", 1)
            patrimoine_mobilier = synthese_data.get('patrimoine_mobilier', {})

            if patrimoine_mobilier:
                doc.add_paragraph(f"Comptes bancaires : {self._safe_get(patrimoine_mobilier, 'comptes_bancaires')}")
                doc.add_paragraph(f"Épargne financière : {self._safe_get(patrimoine_mobilier, 'epargne_financiere')}")
                doc.add_paragraph(f"Assurance-vie : {self._safe_get(patrimoine_mobilier, 'assurance_vie')}")
                doc.add_paragraph(f"Autres investissements : {self._safe_get(patrimoine_mobilier, 'autres_investissements')}")
                doc.add_paragraph(f"Patrimoine mobilier total : {self._safe_get(patrimoine_mobilier, 'patrimoine_mobilier_total')}")
            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === 6. SOCIÉTÉS ===
            self._add_heading_with_style(doc, "6. SOCIÉTÉS", 1)
            societes = synthese_data.get('societes', [])

            if societes:
                for i, societe in enumerate(societes, 1):
                    self._add_heading_with_style(doc, f"6.{i} {self._safe_get(societe, 'raison_sociale')}", 2)
                    doc.add_paragraph(f"Forme juridique : {self._safe_get(societe, 'forme_juridique')}")
                    doc.add_paragraph(f"Pourcentage de détention : {self._safe_get(societe, 'pourcentage_detention')}")
                    doc.add_paragraph(f"Chiffre d'affaires N-1 : {self._safe_get(societe, 'chiffre_affaires_n1')}")
                    doc.add_paragraph(f"Résultat net N-1 : {self._safe_get(societe, 'resultat_net_n1')}")
                    doc.add_paragraph(f"Dettes totales : {self._safe_get(societe, 'dettes_totales')}")
                    doc.add_paragraph(f"Fonds propres : {self._safe_get(societe, 'fonds_propres')}")
                    doc.add_paragraph(f"Activité : {self._safe_get(societe, 'activite')}")
                    doc.add_paragraph()
            else:
                doc.add_paragraph("Non spécifié")
                doc.add_paragraph()

            # === 7. PLAN DE FINANCEMENT ===
            self._add_heading_with_style(doc, "7. PLAN DE FINANCEMENT", 1)
            plan_financement = synthese_data.get('plan_financement', {})

            if plan_financement:
                doc.add_paragraph(f"Apport personnel : {self._safe_get(plan_financement, 'apport_personnel')}")
                doc.add_paragraph(f"Prêt sollicité : {self._safe_get(plan_financement, 'pret_sollicite')}")
                doc.add_paragraph(f"Durée du prêt : {self._safe_get(plan_financement, 'duration_pret')}")
                doc.add_paragraph(f"Taux estimé : {self._safe_get(plan_financement, 'taux_estime')}")
                doc.add_paragraph(f"Mensualité estimée : {self._safe_get(plan_financement, 'mensualite_estimee')}")
                doc.add_paragraph(f"Garanties prévues : {self._safe_get(plan_financement, 'garanties_prevues')}")
                doc.add_paragraph(f"Autres financements : {self._safe_get(plan_financement, 'autres_financements')}")

                # Tableau du financement
                doc.add_paragraph()
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'POSTE'
                hdr_cells[1].text = 'MONTANT'

                # Ajouter les lignes
                postes = [
                    ("Apport personnel", plan_financement.get('apport_personnel', 'Non spécifié')),
                    ("Prêt sollicité", plan_financement.get('pret_sollicite', 'Non spécifié')),
                    ("Autres financements", plan_financement.get('autres_financements', 'Non spécifié'))
                ]

                for poste, montant in postes:
                    row_cells = table.add_row().cells
                    row_cells[0].text = poste
                    row_cells[1].text = montant

            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === 8. ANALYSE FINANCIÈRE ===
            self._add_heading_with_style(doc, "8. ANALYSE FINANCIÈRE", 1)
            analyse_financiere = synthese_data.get('analyse_financiere', {})

            if analyse_financiere:
                doc.add_paragraph(f"Capacité d'emprunt : {self._safe_get(analyse_financiere, 'capacite_emprunt')}")
                doc.add_paragraph(f"Ratio d'endettement : {self._safe_get(analyse_financiere, 'ratio_endettement')}")
                doc.add_paragraph(f"Patrimoine net total : {self._safe_get(analyse_financiere, 'patrimoine_net_total')}")
                doc.add_paragraph(f"Ratio patrimoine/emprunt : {self._safe_get(analyse_financiere, 'ratio_patrimoine_emprunt')}")
                doc.add_paragraph(f"Recommandation : {self._safe_get(analyse_financiere, 'recommandation')}")

                # Points forts
                points_forts = analyse_financiere.get('points_forts', 'Non spécifié')
                if points_forts and points_forts != 'Non spécifié':
                    self._add_heading_with_style(doc, "Points forts", 2)
                    doc.add_paragraph(str(points_forts))

                # Points de vigilance
                points_vigilance = analyse_financiere.get('points_vigilance', 'Non spécifié')
                if points_vigilance and points_vigilance != 'Non spécifié':
                    self._add_heading_with_style(doc, "Points de vigilance", 2)
                    doc.add_paragraph(str(points_vigilance))
            else:
                doc.add_paragraph("Non spécifié")

            doc.add_paragraph()

            # === DOCUMENTS SOURCES ===
            self._add_heading_with_style(doc, "DOCUMENTS SOURCES", 1)
            documents_sources = synthese_data.get('documents_sources', '')
            if documents_sources:
                doc.add_paragraph(documents_sources)

            # === PIED DE PAGE ===
            doc.add_paragraph()
            doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer = doc.add_paragraph("Document généré automatiquement par le système de Carte Financement")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_italic = footer.runs[0]
            footer_italic.italic = True

            # Sauvegarder le document
            doc.save(str(file_path))

            # Enregistrer en base de données
            document_genere_id = insert_document_genere(
                dossier_id=dossier_id,
                type_document="Carte de Financement (.docx)",
                file_path=str(file_path),
                file_name=file_name,
                generated_from=f"Synthèse générée le {synthese_data.get('date_generation', '')}"
            )

            return {
                "success": True,
                "document_genere_id": document_genere_id,
                "file_path": str(file_path),
                "file_name": file_name,
                "dossier_id": dossier_id,
                "message": "Document Word généré avec succès"
            }

        except Exception as e:
            print(f"Erreur lors de la génération du document Word: {e}")
            return {
                "success": False,
                "error": f"Erreur lors de la génération du document Word: {str(e)}"
            }
